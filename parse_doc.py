from docx import Document
import glob, os

desktop = r'C:\Users\老2\Desktop'
files = glob.glob(os.path.join(desktop, '*要让16种*'))
print('Found:', files)

doc = Document(files[0])

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f'[{i}] {para.text}')

print('\n=== TABLES ===')
for ti, table in enumerate(doc.tables):
    print(f'\n--- TABLE {ti} ---')
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n',' | ') for cell in row.cells]
        print(f'  Row {ri}: ' + '  ||  '.join(cells))
